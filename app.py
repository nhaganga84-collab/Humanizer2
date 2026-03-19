import os, sqlite3, io
from datetime import datetime
from pathlib import Path
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, abort
from docx import Document

BASE_DIR = Path(__file__).resolve().parent
DEFAULT_DB_PATH = BASE_DIR / "instance" / "humanizer.db"
FLY_DB_PATH = Path("/data/humanizer.db")
ALLOWED_EXTENSIONS = {"txt", "md", "docx"}

def get_db_path():
    env_path = os.getenv("DATABASE_PATH")
    if env_path:
        return env_path
    if FLY_DB_PATH.parent.exists():
        return str(FLY_DB_PATH)
    return str(DEFAULT_DB_PATH)

def get_connection():
    db_path = Path(get_db_path())
    db_path.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL;")
    return conn

def init_db():
    conn = get_connection()
    conn.execute("CREATE TABLE IF NOT EXISTS history (id INTEGER PRIMARY KEY AUTOINCREMENT, source_type TEXT NOT NULL, original_text TEXT NOT NULL, humanized_text TEXT NOT NULL, tone TEXT NOT NULL, created_at TEXT NOT NULL)")
    conn.commit(); conn.close()

def normalize_spaces(text):
    return "
".join(line.rstrip() for line in text.replace("
", "
").replace("", "
").split("
")).strip()

def split_sentences(text):
    import re
    text = re.sub(r"\s+", " ", text.strip())
    return [] if not text else re.split(r"(?<=[.!?])\s+", text)

def humanize_sentence(sentence, tone):
    reps = {"utilize":"use","facilitate":"help","commence":"start","therefore":"so","moreover":"also","in addition":"also","with regard to":"about","in order to":"to","numerous":"many","individuals":"people","purchase":"buy","obtain":"get","demonstrate":"show","subsequently":"then","prior to":"before","assist":"help"}
    s = sentence.strip()
    for a,b in reps.items():
        s = s.replace(a,b).replace(a.capitalize(), b.capitalize())
    if tone == "casual":
        s = s.replace("do not", "don't").replace("cannot", "can't").replace("I am", "I'm")
    elif tone == "formal":
        s = s.replace("don't", "do not").replace("can't", "cannot")
        if not s.endswith((".","!","?")): s += "."
    elif tone == "academic":
        s = s.replace("help", "support").replace("show", "demonstrate")
        if not s.endswith((".","!","?")): s += "."
    if s and s[0].islower(): s = s[0].upper() + s[1:]
    return s

def humanize_text(text, tone):
    text = normalize_spaces(text)
    if not text: return ""
    out=[]
    for p in [x.strip() for x in text.split("

") if x.strip()]:
        ss = [humanize_sentence(s, tone) for s in split_sentences(p) if s.strip()]
        p2 = " ".join(ss)
        if tone == "casual": p2 = p2.replace("However,", "But").replace("Therefore,", "So")
        elif tone == "formal": p2 = p2.replace("But ", "However, ")
        elif tone == "academic": p2 = p2.replace("But ", "However, ").replace("So ", "Therefore, ")
        out.append(p2)
    return "

".join(out).strip()

def read_docx(file_stream):
    doc = Document(file_stream)
    return "

".join([p.text for p in doc.paragraphs if p.text.strip()])

def build_docx_bytes(title, content):
    d = Document(); d.add_heading(title, level=1)
    for p in content.split("

"):
        if p.strip(): d.add_paragraph(p.strip())
    buf = io.BytesIO(); d.save(buf); buf.seek(0); return buf.read()

def get_recent_history():
    conn = get_connection(); rows = conn.execute("SELECT id, source_type, tone, created_at, substr(original_text, 1, 160) AS preview FROM history ORDER BY id DESC LIMIT 12").fetchall(); conn.close(); return rows

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "dev-secret-key-change-me")
app.config["MAX_CONTENT_LENGTH"] = 8 * 1024 * 1024
init_db()

@app.route("/")
def index():
    return render_template("index.html", history=get_recent_history())

@app.route("/humanize", methods=["POST"])
def humanize():
    tone = request.form.get("tone", "natural")
    text = request.form.get("text", "").strip()
    upload = request.files.get("document")
    source_type = "text"
    if upload and upload.filename:
        ext = upload.filename.rsplit('.',1)[1].lower() if '.' in upload.filename else ''
        if ext not in ALLOWED_EXTENSIONS:
            flash("Formato não suportado. Use .txt, .md ou .docx.")
            return redirect(url_for("index"))
        source_type = ext
        text = read_docx(upload) if ext == "docx" else upload.read().decode("utf-8", errors="ignore")
    if not text.strip():
        flash("Cola um texto ou envia um documento.")
        return redirect(url_for("index"))
    result = humanize_text(text, tone)
    conn = get_connection(); conn.execute("INSERT INTO history (source_type, original_text, humanized_text, tone, created_at) VALUES (?, ?, ?, ?, ?)", (source_type, text, result, tone, datetime.utcnow().isoformat(timespec='seconds'))); conn.commit(); item_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]; conn.close()
    return render_template("index.html", history=get_recent_history(), original_text=text, result_text=result, selected_tone=tone, latest_id=item_id)

@app.route("/history/<int:item_id>")
def history_item(item_id):
    conn = get_connection(); item = conn.execute("SELECT * FROM history WHERE id = ?", (item_id,)).fetchone(); conn.close()
    if not item: abort(404)
    return render_template("index.html", history=get_recent_history(), original_text=item["original_text"], result_text=item["humanized_text"], selected_tone=item["tone"], latest_id=item_id)

@app.route("/download/txt/<int:item_id>")
def download_txt(item_id):
    conn = get_connection(); item = conn.execute("SELECT * FROM history WHERE id = ?", (item_id,)).fetchone(); conn.close()
    if not item: abort(404)
    return send_file(io.BytesIO(item["humanized_text"].encode("utf-8")), as_attachment=True, download_name=f"humanized_{item_id}.txt", mimetype="text/plain; charset=utf-8")

@app.route("/download/docx/<int:item_id>")
def download_docx(item_id):
    conn = get_connection(); item = conn.execute("SELECT * FROM history WHERE id = ?", (item_id,)).fetchone(); conn.close()
    if not item: abort(404)
    return send_file(io.BytesIO(build_docx_bytes("Texto Humanizado", item["humanized_text"])), as_attachment=True, download_name=f"humanized_{item_id}.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/health")
def health():
    return {"status": "ok", "database": get_db_path()}

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", "8080")), debug=True)
