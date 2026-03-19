import os
import re
import sqlite3
import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from flask import Flask, flash,  redirect, render_template, request, send_file, url_for

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = Path('/tmp/humanizer_vercel') if os.getenv('VERCEL') else BASE_DIR
DATA_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = DATA_DIR / 'humanizer.sqlite3'

app = Flask(__name__, template_folder=str(BASE_DIR / 'templates'))
app.secret_key = os.getenv('SECRET_KEY', 'dev-secret-key-change-me')
app.config['MAX_CONTENT_LENGTH'] = 8 * 1024 * 1024

ALLOWED_EXTENSIONS = {'.txt', '.md', '.docx'}
REPLACEMENTS = {
    'utilizar': 'usar',
    'adicionalmente': 'além disso',
    'consequentemente': 'por isso',
    'todavia': 'mas',
    'portanto': 'assim',
    'dessa forma': 'desse jeito',
    'com o objetivo de': 'para',
    'realizar': 'fazer',
    'efetuar': 'fazer',
    'visualizar': 'ver',
    'priorizar': 'dar prioridade a',
    'otimizar': 'melhorar',
    'implementar': 'colocar em prática',
}
OPENERS = ['Na prática,', 'De forma simples,', 'No dia a dia,', 'Em termos claros,']
CLOSERS = [
    'Isso deixa a leitura mais natural.',
    'Assim o texto fica menos mecânico.',
    'Com isso, a mensagem fica mais humana.',
]


def ensure_db() -> None:
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            '''
            CREATE TABLE IF NOT EXISTS humanizations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                created_at TEXT NOT NULL,
                source_type TEXT NOT NULL,
                original_text TEXT NOT NULL,
                humanized_text TEXT NOT NULL,
                tone TEXT NOT NULL,
                notes TEXT
            )
            '''
        )
        conn.commit()


ensure_db()


def get_db_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def normalize_spaces(text: str) -> str:
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def split_sentences(paragraph: str):
    parts = re.split(r'(?<=[.!?])\s+', paragraph.strip())
    return [p.strip() for p in parts if p.strip()]


def vary_sentence(sentence: str, tone: str, index: int) -> str:
    new_sentence = sentence
    for old, new in REPLACEMENTS.items():
        new_sentence = re.sub(rf'\b{re.escape(old)}\b', new, new_sentence, flags=re.IGNORECASE)

    new_sentence = re.sub(r'\b(ChatGPT|IA)\b', 'assistente', new_sentence, flags=re.IGNORECASE)
    new_sentence = re.sub(r'\bimportante ressaltar que\b', 'vale notar que', new_sentence, flags=re.IGNORECASE)
    new_sentence = re.sub(r'\bdevido ao fato de\b', 'porque', new_sentence, flags=re.IGNORECASE)

    if tone == 'casual':
        new_sentence = re.sub(r'\baproximadamente\b', 'mais ou menos', new_sentence, flags=re.IGNORECASE)
        new_sentence = re.sub(r'\bpreviamente\b', 'antes', new_sentence, flags=re.IGNORECASE)
    elif tone == 'formal':
        new_sentence = re.sub(r'\bmas\b', 'porém', new_sentence, flags=re.IGNORECASE)
    elif tone == 'academic':
        new_sentence = re.sub(r'\bpor isso\b', 'por conseguinte', new_sentence, flags=re.IGNORECASE)
        new_sentence = re.sub(r'\bassim\b', 'desse modo', new_sentence, flags=re.IGNORECASE)

    if index == 0 and len(new_sentence.split()) > 8:
        opener = OPENERS[index % len(OPENERS)]
        if not new_sentence.lower().startswith(tuple(o.lower() for o in OPENERS)):
            first = new_sentence[0].lower() + new_sentence[1:] if new_sentence and new_sentence[0].isupper() else new_sentence
            new_sentence = f'{opener} {first}'

    return new_sentence


def humanize_text(text: str, tone: str = 'natural', notes: str = '') -> str:
    text = normalize_spaces(text)
    if not text:
        return ''

    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    output_paragraphs = []

    for p_idx, paragraph in enumerate(paragraphs):
        sentences = split_sentences(paragraph)
        if not sentences:
            continue

        transformed = [vary_sentence(s, tone, i) for i, s in enumerate(sentences)]
        if len(transformed) >= 3:
            transformed[1], transformed[2] = transformed[2], transformed[1]

        if p_idx == len(paragraphs) - 1 and len(transformed[-1].split()) > 6:
            closer = CLOSERS[p_idx % len(CLOSERS)]
            if closer not in transformed[-1]:
                transformed.append(closer)

        output_paragraphs.append(' '.join(transformed))

    result = '\n\n'.join(output_paragraphs)
    if notes.strip():
        result += f'\n\nObservação de estilo aplicada: {notes.strip()}'
    return result


def read_uploaded_file(file_storage) -> str:
    suffix = Path(file_storage.filename).suffix.lower()
    if suffix in {'.txt', '.md'}:
        return file_storage.read().decode('utf-8', errors='ignore')
    if suffix == '.docx':
        doc = Document(BytesIO(file_storage.read()))
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    raise ValueError('Formato não suportado.')


def write_docx(text: str) -> BytesIO:
    doc = Document()
    for paragraph in text.split('\n\n'):
        doc.add_paragraph(paragraph)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _get_history():
    with get_db_connection() as conn:
        return conn.execute(
            'SELECT id, created_at, source_type, tone, substr(original_text, 1, 140) AS preview FROM humanizations ORDER BY id DESC LIMIT 10'
        ).fetchall()


@app.route('/')
def home():
    return render_template("index.html", history=_get_history())


@app.route('/humanize', methods=['POST'])
def humanize():
    tone = request.form.get('tone', 'natural')
    notes = request.form.get('notes', '').strip()
    text = request.form.get('text', '').strip()
    uploaded = request.files.get('document')

    source_type = 'texto'
    if uploaded and uploaded.filename:
        if not allowed_file(uploaded.filename):
            flash('Formato não suportado. Usa .txt, .md ou .docx.', 'error')
            return redirect(url_for('home'))
        text = read_uploaded_file(uploaded)
        source_type = f'documento ({Path(uploaded.filename).suffix.lower()})'

    if not text:
        flash('Escreve um texto ou envia um documento.', 'error')
        return redirect(url_for('home'))

    humanized = humanize_text(text, tone=tone, notes=notes)

    with get_db_connection() as conn:
        conn.execute(
            'INSERT INTO humanizations (created_at, source_type, original_text, humanized_text, tone, notes) VALUES (?, ?, ?, ?, ?, ?)',
            (datetime.utcnow().isoformat(timespec='seconds'), source_type, text, humanized, tone, notes)
        )
        row_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
        conn.commit()

    return render_template(
        'index.html',
        history=_get_history(),
        original_text=text,
        humanized_text=humanized,
        result_id=row_id,
        selected_tone=tone,
        notes=notes,
    )


@app.route('/history/<int:item_id>')
def history_item(item_id: int):
    with get_db_connection() as conn:
        item = conn.execute('SELECT * FROM humanizations WHERE id = ?', (item_id,)).fetchone()
    if not item:
        flash('Registo não encontrado.', 'error')
        return redirect(url_for('home'))

    return render_template(
        'index.html',
        history=_get_history(),
        original_text=item['original_text'],
        humanized_text=item['humanized_text'],
        result_id=item['id'],
        selected_tone=item['tone'],
        notes=item['notes'] or '',
    )


@app.route('/download/<int:item_id>/<fmt>')
def download(item_id: int, fmt: str):
    with get_db_connection() as conn:
        item = conn.execute('SELECT * FROM humanizations WHERE id = ?', (item_id,)).fetchone()
    if not item:
        flash('Registo não encontrado.', 'error')
        return redirect(url_for('home'))

    safe_name = f'humanizado_{item_id}_{uuid.uuid4().hex[:6]}'
    if fmt == 'txt':
        data = BytesIO(item['humanized_text'].encode('utf-8'))
        return send_file(data, as_attachment=True, download_name=f'{safe_name}.txt', mimetype='text/plain; charset=utf-8')
    if fmt == 'docx':
        return send_file(write_docx(item['humanized_text']), as_attachment=True, download_name=f'{safe_name}.docx')

    flash('Formato de download inválido.', 'error')
    return redirect(url_for('history_item', item_id=item_id))


@app.route('/delete/<int:item_id>', methods=['POST'])
def delete_item(item_id: int):
    with get_db_connection() as conn:
        conn.execute('DELETE FROM humanizations WHERE id = ?', (item_id,))
        conn.commit()
    flash('Registo removido.', 'success')
    return redirect(url_for('home'))
